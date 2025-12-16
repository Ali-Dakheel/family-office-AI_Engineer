"""
ADOR - Augmented Document Reader
DOCX Entity Extractor Module

This module implements a rule-based parser for extracting financial entities
from structured DOCX documents. It leverages document structure (tables, paragraphs)
and pattern matching to achieve deterministic, high-accuracy extraction.
"""

import re
import json
from pathlib import Path
from typing import Dict, List, Optional, Any
from dataclasses import dataclass, asdict
from datetime import datetime
from docx import Document


@dataclass
class FinancialEntity:
    """Data class representing extracted financial entities."""
    counterparty: Optional[str] = None
    initial_valuation_date: Optional[str] = None
    notional: Optional[str] = None
    valuation_date: Optional[str] = None
    maturity: Optional[str] = None
    underlying: Optional[str] = None
    coupon: Optional[str] = None
    barrier: Optional[str] = None
    calendar: Optional[str] = None
    
    # Metadata
    extraction_timestamp: str = None
    source_document: str = None
    
    def __post_init__(self):
        """Set extraction timestamp if not provided."""
        if self.extraction_timestamp is None:
            self.extraction_timestamp = datetime.utcnow().isoformat()
    
    def to_dict(self) -> Dict[str, Any]:
        """Convert entity to dictionary, excluding None values."""
        return {k: v for k, v in asdict(self).items() if v is not None}


class DOCXEntityExtractor:
    """
    Rule-based entity extractor for DOCX financial documents.
    
    This extractor uses a multi-strategy approach:
    1. Table-based extraction for structured data
    2. Paragraph pattern matching for header information
    3. Regex-based cleaning and normalization
    """
    
    # Entity field mappings (table row label → entity field)
    ENTITY_MAPPINGS = {
        'party a': 'counterparty',
        'counterparty': 'counterparty',
        'initial valuation date': 'initial_valuation_date',
        'notional': 'notional',
        'notional amount': 'notional',
        'notional amount (n)': 'notional',
        'valuation date': 'valuation_date',
        'maturity': 'maturity',
        'termination date': 'maturity',
        'maturity date': 'maturity',
        'underlying': 'underlying',
        'coupon': 'coupon',
        'coupon (c)': 'coupon',
        'barrier': 'barrier',
        'barrier (b)': 'barrier',
        'calendar': 'calendar',
        'business day': 'calendar',
    }
    
    def __init__(self, document_path: str):
        """
        Initialize the extractor with a DOCX document.
        
        Args:
            document_path: Path to the DOCX file
        """
        self.document_path = Path(document_path)
        if not self.document_path.exists():
            raise FileNotFoundError(f"Document not found: {document_path}")
        
        self.doc = Document(str(self.document_path))
        self.entity = FinancialEntity(source_document=self.document_path.name)
    
    def extract(self) -> FinancialEntity:
        """
        Main extraction method orchestrating all extraction strategies.
        
        Returns:
            FinancialEntity object with extracted data
        """
        # Strategy 1: Extract from paragraphs (header information)
        self._extract_from_paragraphs()
        
        # Strategy 2: Extract from tables (structured data)
        self._extract_from_tables()
        
        # Strategy 3: Post-processing and validation
        self._post_process()
        
        return self.entity
    
    def _extract_from_paragraphs(self) -> None:
        """Extract entities from document paragraphs using pattern matching."""
        for para in self.doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            
            # Pattern: "Barrier XX%, DD Month YYYY"
            barrier_match = re.search(r'Barrier\s+(\d+(?:\.\d+)?%)', text, re.IGNORECASE)
            if barrier_match and not self.entity.barrier:
                self.entity.barrier = barrier_match.group(1)
            
            # Extract date from barrier line (if present)
            date_match = re.search(r'(\d{1,2}\s+\w+\s+\d{4})', text)
            if date_match and barrier_match and not self.entity.maturity:
                self.entity.maturity = date_match.group(1)
    
    def _extract_from_tables(self) -> None:
        """Extract entities from document tables using field mappings."""
        for table in self.doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                
                # Skip empty rows or rows with less than 2 cells
                if len(cells) < 2 or not cells[0]:
                    continue
                
                # Normalize the field label
                field_label = cells[0].lower().strip()
                field_value = cells[1].strip()
                
                # Skip if value is empty or placeholder
                if not field_value or field_value.startswith('***'):
                    continue
                
                # Map to entity field
                entity_field = self.ENTITY_MAPPINGS.get(field_label)
                if entity_field:
                    # Only set if not already set (first occurrence takes precedence)
                    current_value = getattr(self.entity, entity_field)
                    if current_value is None:
                        setattr(self.entity, entity_field, field_value)
    
    def _post_process(self) -> None:
        """Post-process and normalize extracted entities."""
        # Clean and normalize coupon values
        if self.entity.coupon:
            self.entity.coupon = self.entity.coupon.strip()
        
        # Normalize barrier format
        if self.entity.barrier:
            # Extract percentage if embedded in longer text
            barrier_match = re.search(r'(\d+(?:\.\d+)?%)', self.entity.barrier)
            if barrier_match:
                self.entity.barrier = barrier_match.group(1)
        
        # Normalize notional format
        if self.entity.notional:
            # Keep full format with currency
            self.entity.notional = re.sub(r'\s+', ' ', self.entity.notional).strip()
    
    def export_json(self, output_path: Optional[str] = None) -> str:
        """
        Export extracted entities to JSON format.
        
        Args:
            output_path: Optional path for output file. If None, generates default path.
        
        Returns:
            Path to the output JSON file
        """
        if output_path is None:
            output_dir = self.document_path.parent.parent / 'outputs'
            output_dir.mkdir(exist_ok=True)
            output_path = output_dir / f"{self.document_path.stem}_entities.json"
        
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        
        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(self.entity.to_dict(), f, indent=2, ensure_ascii=False)
        
        return str(output_path)
    
    def export_summary(self) -> str:
        """Generate a human-readable summary of extracted entities."""
        lines = ["=" * 60, "EXTRACTED FINANCIAL ENTITIES", "=" * 60, ""]
        
        for field, value in self.entity.to_dict().items():
            if field not in ['extraction_timestamp', 'source_document']:
                label = field.replace('_', ' ').title()
                lines.append(f"{label:.<30} {value}")
        
        lines.extend(["", "=" * 60, 
                     f"Source: {self.entity.source_document}",
                     f"Extracted: {self.entity.extraction_timestamp}",
                     "=" * 60])
        
        return "\n".join(lines)


def main():
    """Main execution function for standalone usage."""
    import argparse
    
    parser = argparse.ArgumentParser(
        description='Extract financial entities from DOCX documents'
    )
    parser.add_argument(
        'input_file',
        help='Path to input DOCX file'
    )
    parser.add_argument(
        '-o', '--output',
        help='Output JSON file path (optional)',
        default=None
    )
    parser.add_argument(
        '-v', '--verbose',
        action='store_true',
        help='Print extraction summary'
    )
    
    args = parser.parse_args()
    
    try:
        # Extract entities
        extractor = DOCXEntityExtractor(args.input_file)
        entity = extractor.extract()
        
        # Export to JSON
        output_path = extractor.export_json(args.output)
        
        if args.verbose:
            print(extractor.export_summary())
            print(f"\nResults saved to: {output_path}")
        else:
            print(f"Extraction complete: {output_path}")
        
        return 0
    
    except Exception as e:
        print(f"Error: {e}", file=__import__('sys').stderr)
        return 1


if __name__ == '__main__':
    exit(main())